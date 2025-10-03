import os
import tempfile

import pandas as pd
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader

# Page configuration
st.set_page_config(page_title="CV Sorter", layout="wide")


def main():
    st.title("📄 Intelligent CV Sorting System")

    # Sidebar for settings
    with st.sidebar:
        st.header("Configuration")

        # OpenAI API Key
        openai_api_key = st.text_input("OpenAI API Key", type="password")

        # Job description
        job_description = st.text_area(
            "Job Description",
            height=200,
            placeholder="Paste the job description here...",
        )

        # Job description file upload
        uploaded_jd_file = st.file_uploader(
            "Or upload a job description", type=["txt", "pdf", "docx"]
        )

        if uploaded_jd_file:
            try:
                if uploaded_jd_file.type == "application/pdf":
                    reader = PdfReader(uploaded_jd_file)
                    job_description = " ".join(
                        [
                            page.extract_text()
                            for page in reader.pages
                            if page.extract_text()
                        ]
                    )

                elif (
                    uploaded_jd_file.type
                    == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ):
                    doc = Document(uploaded_jd_file)
                    job_description = " ".join(
                        [para.text for para in doc.paragraphs if para.text]
                    )

                else:
                    job_description = str(uploaded_jd_file.read(), "utf-8")

                st.success("Job description loaded successfully!")

            except Exception as e:
                st.error(f"Error reading file: {e}")

        # CV uploads
        uploaded_cvs = st.file_uploader(
            "Upload CVs (PDF/DOCX)",
            accept_multiple_files=True,
            type=["pdf", "docx"],
        )

    # Main section
    if openai_api_key and job_description and uploaded_cvs:

        # Function to extract text from a file
        def extract_text_from_file(file_path, file_type):
            try:
                if file_type == "pdf":
                    reader = PdfReader(file_path)
                    return " ".join(
                        [
                            page.extract_text()
                            for page in reader.pages
                            if page.extract_text()
                        ]
                    )

                elif file_type == "docx":
                    doc = Document(file_path)
                    return " ".join([para.text for para in doc.paragraphs if para.text])

            except Exception as e:
                st.error(f"Error reading {file_path}: {e}")
                return ""
            pass

        # CV processing
        if st.button("🚀 Analyze CVs"):

            results = []

            with st.spinner("📊 Analyzing CVs in progress..."):

                for uploaded_file in uploaded_cvs:
                    try:
                        # Temporarily save the file
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=os.path.splitext(uploaded_file.name)[1]
                        ) as tmp_file:
                            tmp_file.write(uploaded_file.read())
                            tmp_path = tmp_file.name

                        # Determine file type
                        file_type = (
                            "pdf" if uploaded_file.type == "application/pdf" else "docx"
                        )

                        # Extract text
                        cv_text = extract_text_from_file(tmp_path, file_type)

                        if cv_text:
                            # Evaluation simulation (to be replaced by OpenAI)
                            score = len(cv_text) % 100  # Simple simulation
                            match_percent = min(score + 20, 100)  # Simulation

                            results.append(
                                {
                                    "File": uploaded_file.name,
                                    "Score": score,
                                    "Match %": match_percent,
                                    "Strengths": (
                                        "Relevant experience"
                                        if score > 50
                                        else "Basic skills"
                                    ),
                                    "Weaknesses": (
                                        "Lack of experience"
                                        if score <= 50
                                        else "Limited info"
                                    ),
                                }
                            )

                        # Clean up temporary file
                        os.unlink(tmp_path)

                    except Exception as e:
                        st.error(f"Error with {uploaded_file.name}: {e}")

            if results:
                # Create and display results table
                results_df = pd.DataFrame(results).sort_values("Score", ascending=False)

                st.subheader("🎯 Analysis Results")
                st.dataframe(results_df, use_container_width=True)

                # Metrics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric(
                        "📊 Average Score", f"{results_df['Score'].mean():.1f}/100"
                    )
                with col2:
                    st.metric("⭐ Best Score", f"{results_df['Score'].max()}/100")
                with col3:
                    st.metric("📋 CVs Analyzed", len(results_df))

                # Chart
                st.subheader("📈 CV Performance")
                st.bar_chart(results_df.set_index("File")["Score"])

                # Download results
                csv = results_df.to_csv(index=False).encode("utf-8")
                st.download_button(
                    "💾 Download Results (CSV)",
                    csv,
                    "cv_analysis_results.csv",
                    "text/csv",
                    key="download-csv",
                )
            else:
                st.warning("❌ No CVs could be analyzed")

    else:
        # Instructions message
        st.info(
            """
        ## 📋 Instructions
        1. Enter your OpenAI API key
        2. Enter or upload a job description
        3. Upload CVs to analyze (PDF or DOCX)
        4. Click 'Analyze CVs'
        """
        )

        if not openai_api_key:
            st.warning("🔑 OpenAI API key required")
        if not job_description:
            st.warning("📝 Job description required")
        if not uploaded_cvs:
            st.warning("📄 CVs to analyze required")


# Entry point
if __name__ == "__main__":
    main()
